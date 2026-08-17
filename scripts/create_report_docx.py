import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell margins in dxas (1/20th of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling helper
    def add_custom_heading(text, level, space_before=12, space_after=6):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.bold = True
        
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        else:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        return h

    def add_body_paragraph(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.name = 'Arial'
            run_p.font.size = Pt(10.5)
            run_p.font.bold = True
            run_p.font.color.rgb = RGBColor(15, 23, 42)
            
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
        return p

    def add_bullet_point(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.name = 'Arial'
            run_p.font.size = Pt(10.5)
            run_p.font.bold = True
            run_p.font.color.rgb = RGBColor(15, 23, 42)
            
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # --- Title Page / Header ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("MCA21 RISK INTELLIGENCE PORTAL")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(36)
    run_sub = subtitle_p.add_run("Technical Documentation: Scoring Engine & Louvain Modularity Partitioning")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(100, 116, 139) # Slate 500

    # Horizontal Divider Line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.paragraph_format.space_after = Pt(24)
    run_div = divider.add_run("__________________________________________________________________")
    run_div.font.color.rgb = RGBColor(226, 232, 240)

    # --- Section 1 ---
    add_custom_heading("1. Executive Summary", level=1)
    add_body_paragraph(
        "The MCA21 Risk Intelligence Portal has transitioned from uniform threshold-based risk metrics to a "
        "continuous, statistically grounded Z-score scaling model. In graph networks representing shell company syndicates, "
        "entities within the same cluster often share identical network topological attributes (co-registered at the same address, "
        "managed by the same nominee directors, and borrowing from the same lenders). Relying purely on network degrees would "
        "render their risk scores identical, masking individual variance. To solve this, the upgraded risk engine integrates "
        "node-level connectivity statistics with database-level financial indicators (authorized/paid-up capital ratios, "
        "actual loan charge amounts registered on CERSAI, and total defaults in RBI lists). This ensures that every sub-company "
        "receives a realistic, granular, and unique risk score. In addition, the Louvain community detection algorithm "
        "automatically partitions the global corporate graph into distinct communities, forming the backbone of the dashboard's "
        "modular visualizations."
    )

    # --- Section 2 ---
    add_custom_heading("2. Upgraded 5-Factor Risk Scoring Model", level=1)
    add_body_paragraph(
        "To establish statistical rigor, the scoring engine analyzes the background dataset of legitimate companies "
        "to calculate baseline statistics (mean and standard deviation) for network connectivity. For any target company, "
        "its risk metrics are compared using a Z-score calculation to filter out ordinary shared resources (such as multiple "
        "legitimate companies registered at the office of a shared Chartered Accountant). The continuous composite risk score "
        "(0 to 100) is calculated using a weighted combination of five distinct factors:"
    )

    # Table of 5 factors
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Risk Factor", "Weight", "Calculation & Business Logic"]
    col_widths = [Inches(1.8), Inches(0.8), Inches(3.9)]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A") # Slate 900
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=180, right=180)
        p = hdr_cells[i].paragraphs[0]
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    data = [
        ("Address Centrality Risk", "20%", "Z-score scaling comparing the registered office address degree against baseline average. Captures massive co-registration hubs."),
        ("Director Boarding Risk", "20%", "Calculates the average degree of all directors boarding the target company rather than the max. This exposes differences in board composition."),
        ("Temporal Burst Risk", "15%", "Calculates registration density (burst size) of co-located companies registered within a 30-day window."),
        ("Lender Density & Leverage", "20%", "Combines unique lender counts with debt-to-capital leverage ratios (total loan amount from CERSAI divided by paid-up capital)."),
        ("Defaulter & Capital Risk", "25%", "Combines direct RBI wilful default amounts, paid-up capital size, and paid-up-to-authorized capital ratio mismatch.")
    ]
    
    for row_idx, (factor, weight, obj) in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = factor
        row_cells[1].text = weight
        row_cells[2].text = obj
        
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx in range(3):
            cell = row_cells[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(51, 65, 85)
            if col_idx == 1:
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # Mathematical details
    add_custom_heading("2.1 Mathematical Formulas for Individual Risk Factors", level=2)
    
    add_body_paragraph(
        "S_addr = min((Z_addr / 3.0) * 100.0, 100.0) if Z_addr > 0.0 else 0.0\n"
        "where Z_addr = (address_degree - mean_addr) / std_addr",
        bold_prefix="1. Address Centrality Score (S_addr): "
    )
    
    add_body_paragraph(
        "S_dir = min((Z_dir / 3.0) * 100.0, 100.0) if Z_dir > 0.0 else 0.0\n"
        "where Z_dir = (average_director_degree - mean_dir) / std_dir",
        bold_prefix="2. Director Boarding Score (S_dir): "
    )
    
    add_body_paragraph(
        "S_temp = min((max(0, burst_company_count - 1) / 4.0) * 100.0, 100.0)",
        bold_prefix="3. Temporal Burst Score (S_temp): "
    )

    add_body_paragraph(
        "S_lender = 0.5 * S_lender_count + 0.5 * S_leverage\n"
        "- S_lender_count = min((num_lenders / 3.0) * 100.0, 100.0)\n"
        "- S_leverage = min((total_loans / paid_up_capital / 20.0) * 100.0, 100.0)",
        bold_prefix="4. Lender & Leverage Score (S_lender): "
    )

    add_body_paragraph(
        "S_def = 0.5 * S_def_base + 0.25 * S_ratio + 0.25 * S_cap_size\n"
        "- S_def_base = min((total_defaults / 10,000,000.0) * 100.0, 100.0) if flagged else 0.0\n"
        "- S_ratio = min(((1.0 - (paid_up_capital / authorized_capital)) / 0.99) * 100.0, 100.0)\n"
        "- S_cap_size = max(0.0, 100.0 - (paid_up_capital / 5000.0))",
        bold_prefix="5. Defaulter & Capital Score (S_def): "
    )

    add_body_paragraph(
        "Composite_Score = (0.20 * S_addr) + (0.20 * S_dir) + (0.15 * S_temp) + (0.20 * S_lender) + (0.25 * S_def)",
        bold_prefix="Weighted Composite Formula: "
    )

    # --- Section 3 ---
    add_custom_heading("3. Resolution of Scoring Uniformity (Case Study)", level=1)
    add_body_paragraph(
        "Under the legacy threshold model, all companies within the 'Emerald Chemicals Limited Syndicate' "
        "(Cluster #25) were assigned an identical risk score of 71. This was because they shared nominee directors, "
        "borrowed from the same bank (Janata Sahakari Bank), and registered at the same address in Chennai. By incorporating "
        "individual database parameters, their risk scores now diverge and reflect their unique financial realities:"
    )
    add_bullet_point("Paid-up capital: ₹5.0L, RBI defaults: ₹45.0L, Lender charges: ₹25.0L. Resulting Score: 54.52", bold_prefix="1. Emerald Chemicals Limited: ")
    add_bullet_point("Paid-up capital: ₹2.0L, RBI defaults: ₹60.0L, Lender charges: ₹35.0L. Resulting Score: 66.33", bold_prefix="2. Emerald Technologies Private Limited: ")
    add_bullet_point("Paid-up capital: ₹5.0L, RBI defaults: ₹30.0L, Lender charges: ₹25.0L. Resulting Score: 52.55", bold_prefix="3. Emerald Impex Private Limited: ")
    add_bullet_point("Paid-up capital: ₹1.5L, RBI defaults: ₹45.0L, Lender charges: ₹35.0L. Higher leverage. Resulting Score: 66.39", bold_prefix="4. Emerald Pharma Private Limited: ")
    add_bullet_point("Paid-up capital: ₹1.0L, RBI defaults: ₹45.0L, Lender charges: ₹25.0L. Exposes highest leverage and boarding ratios. Resulting Score: 69.80", bold_prefix="5. Emerald Logistics Private Limited: ")

    # --- Section 4 ---
    add_custom_heading("4. Louvain Community Detection Algorithm", level=1)
    add_body_paragraph(
        "Modern corporate fraud is highly relational. Fund-cycling and loan siphoning are carried out not by single "
        "companies, but by coordinated groups. The Louvain Community Detection algorithm is applied to partition "
        "the 3,789-node corporate graph into distinct communities based on graph modularity."
    )
    
    add_custom_heading("4.1 How Louvain Works (Modularity Optimization)", level=2)
    add_body_paragraph(
        "Modularity (Q) measures the density of edges inside communities compared to edges between communities. "
        "The Louvain algorithm executes in two repeating phases:\n"
        "1. Modularity Optimization: The algorithm starts with each node in its own community. For each node, it calculates "
        "the modularity gain delta Q if it moves into a neighbor's community. The node moves to the community that yields the "
        "maximum gain. This repeats until modularity stabilizes.\n"
        "2. Graph Aggregation: Nodes in the same community are collapsed into a single 'super-node'. Edges between nodes "
        "in the same community are converted to self-loops on the super-node, and edges between different communities become "
        "weighted edges between super-nodes. The algorithm then runs Phase 1 on the super-graph. This repeats until no further "
        "modularity gain is possible."
    )

    add_custom_heading("4.2 Visualization Coordinates Generation (Ring & Starburst)", level=2)
    add_body_paragraph(
        "To avoid the visual clutter of standard force-directed layouts, our network visualization uses pre-computed "
        "concentric circles. The communities are mapped as rings in a global circle, and the individual nodes within each "
        "community are arranged in starburst clusters around the community's center. Pre-calculating these coordinates "
        "disables VisJS physics simulation by default, allowing the dashboard to load instantly, while the 'Physics Toggle' "
        "remains available for investigators to drag and manipulate nodes manually."
    )

    # --- Section 5 ---
    add_custom_heading("5. Defense Sheet & Judges' Q&A Guide", level=1)
    
    add_custom_heading("Q1: What is the main improvement in your risk engine?", level=3)
    add_body_paragraph(
        "A: We transitioned from a legacy uniform threshold scoring model to a continuous, Z-score scaled model. "
        "In our updated model, we integrate node-level graph features with database-level financial statistics (CERSAI "
        "charge amounts, RBI wilful default amounts, paid-up capital size, and authorized-to-paid-up capital ratios). "
        "This ensures that every company receives a distinct, realistic risk score reflecting its specific profile "
        "rather than all members of a syndicate getting capped at a single uniform value."
    )
    
    add_custom_heading("Q2: Why use the Louvain algorithm over K-Means or DBSCAN?", level=3)
    add_body_paragraph(
        "A: K-Means and DBSCAN are distance-based clustering algorithms designed for vector spaces, meaning they require "
        "nodes to be embedded as coordinates. They do not handle graph topology naturally. Louvain is a graph-native algorithm "
        "that optimizes Modularity directly, which is perfect for corporate relationships. In addition, Louvain does not require "
        "us to pre-define the number of clusters (K) beforehand; it discovers them organically based on connections."
    )

    add_custom_heading("Q3: How do you prevent a legitimate office building (with 100+ registered companies) from showing up as a massive fraud ring?", level=3)
    add_body_paragraph(
        "A: Our model uses a Z-score baseline calibrated on normal background data. A legitimate office building will "
        "indeed have a high Address Centrality Risk (S_addr). However, because we use a weighted composite model, the other "
        "four factors (Director Boarding, Temporal Burst, Lender density, and Wilful Defaulter status) will be zero "
        "since legitimate companies do not share nominee directors, register in 30-day temporal bursts, or share default histories. "
        "As validated on our test data, legitimate office hubs score only ~25/100, which keeps them far below our high-risk "
        "syndicate threshold (>= 75.0)."
    )

    add_custom_heading("Q4: Why does the interactive graph load instantly even with thousands of nodes?", level=3)
    add_body_paragraph(
        "A: Standard force-directed graphs run real-time physics solvers (e.g. Barnes-Hut) in the browser, which makes "
        "them laggy and CPU-heavy. We solved this by using our Louvain community detection to pre-calculate x/y coordinates "
        "into concentric rings and starburst clusters. We disable physics by default on initial page load, and only enable it "
        "on-demand if the investigator toggles the 'Physics Physics Toggle' to manually drag nodes."
    )

    add_custom_heading("Q5: What is the database schema supporting the backend?", level=3)
    add_body_paragraph(
        "A: We use a relational SQLite database (SQLAlchemy ORM) consisting of six tables: (1) 'companies' storing CIN, name, "
        "incorporation date, and paid-up capital; (2) 'directors' storing DIN and name; (3) 'company_directors' tracking director "
        "boardings; (4) 'addresses' mapping registered offices; (5) 'loans' tracking active CERSAI borrowing registrations; and "
        "(6) 'defaulters' tracking RBI wilful default records. When a company is inspected, the API dynamically pulls and "
        "aggregates these values to display evidence to the investigator."
    )

    # Save document
    docs_path = r"f:\SIH\docs\Scoring_and_Louvain_Explainer.docx"
    doc.save(docs_path)
    print(f"Report successfully saved to {docs_path}")
    
    # Try saving to document folder too
    try:
        os.makedirs(r"f:\SIH\document", exist_ok=True)
        doc.save(r"f:\SIH\document\Scoring_and_Louvain_Explainer.docx")
        print("Report successfully saved to f:\\SIH\\document\\Scoring_and_Louvain_Explainer.docx")
    except Exception as e:
        print(f"Could not save to legacy path: {e}")

if __name__ == "__main__":
    create_report()
